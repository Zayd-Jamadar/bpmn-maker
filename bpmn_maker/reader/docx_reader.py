from docx import Document

from .base import IDocumentReader, RawDocument


class DocxReader(IDocumentReader):
    """
    Reads a .docx file using python-docx.
    Returns every non-empty paragraph as a string in document order.
    Strips leading/trailing whitespace; ignores table cells for now.
    """

    def read(self, path: str) -> RawDocument:
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        metadata = {"title": doc.core_properties.title or ""}
        return RawDocument(paragraphs=paragraphs, metadata=metadata)
